#!/usr/bin/env python3
"""生成校内审批一页纸说明（Word + PDF）。"""

from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "school-approval-application.docx"
OUT_PDF = ROOT / "docs" / "school-approval-application.pdf"

# 可选：从 frontend/src/config/app.js 同步（简单解析占位）
CONFIG = ROOT / "frontend" / "src" / "config" / "app.js"


def read_config_defaults():
    org = "__________________________ 学院"
    email = "__________________________"
    advisor = "__________"
    if CONFIG.exists():
        text = CONFIG.read_text(encoding="utf-8")
        for key, var in [
            ("orgName", org),
            ("contactEmail", email),
            ("advisorName", advisor),
        ]:
            marker = f"{key}: '"
            if marker in text:
                start = text.index(marker) + len(marker)
                end = text.index("'", start)
                val = text[start:end]
                if "待填写" not in val and "example" not in val:
                    if key == "orgName":
                        org = val
                    elif key == "contactEmail":
                        email = val
                    elif key == "advisorName":
                        advisor = val
    return org, email, advisor


def build_docx(org, email, advisor):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("关于申请以学校组织主体注册\n「西语背单词」微信小程序的说明")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph("（一页纸 · 提交指导教师 / 创新创业学院 / 网信中心）")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "西语背单词（DELE 分级识记）"),
        ("申请单位", org),
        ("指导教师", f"{advisor}　职称：__________　电话：__________"),
        ("学生负责人", "__________　学号：__________　电话：__________"),
        ("项目联系邮箱", email),
        ("申请日期", "________ 年 ____ 月 ____ 日"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    sections = [
        (
            "一、申请事项",
            "恳请学校/院系同意以学校组织主体（事业单位法人）在微信公众平台注册并上线微信小程序"
            "「西语背单词」，用于本校西班牙语专业学生的 DELE 分级词汇学习，并支撑大学生创新创业"
            "训练计划（大创）项目申报与成果展示。\n"
            "建议小程序管理员由指导教师担任；学生团队负责开发、内容维护与日常运营协助。",
        ),
        (
            "二、项目简介",
            "「西语背单词」是一款面向 DELE 考试体系的西班牙语识记工具，采用「配图 + 四选一」"
            "学习模式，并集成错题本、学习统计、间隔复习（SM-2）等功能。词库与例句由西班牙语"
            "专业同学编制与校对，技术由项目组自主开发（uni-app + 后端 API）。\n"
            "预期用户：本校及同类院校西语专业学生、DELE 备考者。\n"
            "当前进度：H5 可交互原型与本地后端已完成；隐私政策、用户协议及首次隐私同意流程"
            "已实现，待主体与域名就绪后可提交微信审核。",
        ),
        (
            "三、合规与安全承诺",
            "· 是否收费：否（无支付、无广告）\n"
            "· 用户生成内容（UGC）：无（用户仅学习，不上传文字/图片）\n"
            "· 数据采集：仅微信 openid 与学习进度；不采集手机号、位置、通讯录\n"
            "· 数据传输与存储：全站 HTTPS；数据存储于境内服务器\n"
            "· 内容表述：不使用「官方 DELE」「保证过级」等未授权或绝对化用语\n"
            "· 知识产权：词库、配图为团队自制或合法授权，已建立版权台账",
        ),
        (
            "四、所需学校支持",
            "1. 主体与证照：同意以学校/院系组织主体注册小程序，并提供组织机构证照等材料。\n"
            "2. 管理员与认证：指导教师担任小程序管理员；微信认证费用约 300 元/年"
            "（经费来源：__________）。\n"
            "3. 备案与域名：协助或指引完成小程序 ICP 备案及 API 域名备案"
            "（计划域名：__________）。\n"
            "4. 法务/网信审阅（若需）：审阅《用户协议》《隐私政策》定稿后上线。",
        ),
        (
            "五、上线计划（约 6 周）",
            "第 1 周：院系审批、主体申请意向确认\n"
            "第 2–3 周：域名备案、法务/网信材料\n"
            "第 4 周：微信注册、认证、小程序备案\n"
            "第 5 周：生产环境部署、体验版验收\n"
            "第 6 周：提交审核、发布正式版",
        ),
        (
            "六、附件清单",
            "□ 立项书摘要 / 答辩 PPT\n"
            "□ 功能截图（首页、学习、错题本、统计）\n"
            "□ 团队成员及分工说明\n"
            "□ 《用户协议》《隐私政策》草案\n"
            "□ 配图版权台账说明",
        ),
        (
            "七、审批意见（院系 / 创院 / 网信填写）",
            "□ 同意　以学校组织主体注册上述小程序，并按学校信息化与合规要求推进备案与上线。\n"
            "□ 不同意　原因：________________________________________________\n"
            "□ 需补充材料　________________________________________________\n\n"
            "指导教师签字：________________　日期：________________\n"
            "院系负责人签字：________________　日期：________________\n"
            "创新创业学院（如适用）签字：________________　日期：________________\n"
            "网信中心（如适用）签字：________________　日期：________________\n\n"
            "学生负责人签字：________________　日期：________________",
        ),
    ]

    for heading, body in sections:
        h = doc.add_paragraph()
        h.add_run(heading).bold = True
        for para in body.split("\n"):
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10.5)

    foot = doc.add_paragraph(
        "本说明为项目组起草模板，下划线处请填写后提交。详细核对表见《校内审批材料清单》。"
    )
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].italic = True

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


def build_pdf(org, email, advisor):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_paths = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    font_name = "Helvetica"
    for fp in font_paths:
        if Path(fp).exists():
            try:
                pdfmetrics.registerFont(TTFont("CJK", fp, subfontIndex=0))
                font_name = "CJK"
                break
            except Exception:
                continue

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5,
        leading=14,
        spaceAfter=4,
    )
    h_style = ParagraphStyle(
        "HCN",
        parent=body_style,
        fontSize=10.5,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.black,
    )

    story = []
    story.append(Paragraph("关于申请以学校组织主体注册<br/>「西语背单词」微信小程序的说明", title_style))
    story.append(Paragraph("（一页纸 · 提交指导教师 / 创新创业学院 / 网信中心）", body_style))
    story.append(Spacer(1, 0.3 * cm))

    meta = [
        ["项目名称", "西语背单词（DELE 分级识记）"],
        ["申请单位", org],
        ["指导教师", f"{advisor}　职称：__________　电话：__________"],
        ["学生负责人", "__________　学号：__________　电话：__________"],
        ["项目联系邮箱", email],
        ["申请日期", "________ 年 ____ 月 ____ 日"],
    ]
    t = Table(meta, colWidths=[3.2 * cm, 12.8 * cm])
    t.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font_name, 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f5f5f5")),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 0.2 * cm))

    blocks = [
        (
            "一、申请事项",
            "恳请学校/院系同意以<strong>学校组织主体</strong>在微信公众平台注册并上线微信小程序"
            "「西语背单词」，用于 DELE 分级词汇学习及大创项目展示。"
            "建议管理员由指导教师担任，学生团队负责开发与运营协助。",
        ),
        (
            "二、项目简介",
            "面向 DELE 的西班牙语识记工具：配图 + 四选一、错题本、统计、SM-2 复习。"
            "词库由西语专业同学编制，技术自主开发。H5 原型与隐私合规流程已完成。",
        ),
        (
            "三、合规承诺",
            "不收费、无 UGC；仅采集 openid 与学习记录；HTTPS + 境内存储；"
            "不使用「官方 DELE」「保证过级」等表述。",
        ),
        (
            "四、所需支持",
            "组织主体与证照；教师任管理员（认证约 300 元/年）；备案与域名协助；"
            "协议审阅（若需）。",
        ),
        (
            "五、审批意见",
            "□ 同意　□ 不同意　□ 需补充材料<br/>"
            "指导教师 / 院系 / 创院 / 网信签字栏（见 Word 完整版）<br/>"
            "学生负责人签字：________________　日期：________________",
        ),
    ]
    for head, text in blocks:
        story.append(Paragraph(f"<b>{head}</b>", h_style))
        story.append(Paragraph(text, body_style))

    story.append(Spacer(1, 0.2 * cm))
    story.append(
        Paragraph(
            "<i>PDF 为精简版；完整表格与签字栏请使用同目录 Word 文件。</i>",
            body_style,
        )
    )

    doc.build(story)
    print(f"Wrote {OUT_PDF}")


def main():
    org, email, advisor = read_config_defaults()
    build_docx(org, email, advisor)
    build_pdf(org, email, advisor)


if __name__ == "__main__":
    main()
